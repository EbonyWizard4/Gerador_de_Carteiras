import pandas as pd

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.enum.table import WD_TABLE_ALIGNMENT

def Do_docx():
    document = Document()

    marcaTxt = 'Ebw invest'
    marca = document.add_paragraph(marcaTxt, style='Intense Quote')
    marca.alignment =  WD_PARAGRAPH_ALIGNMENT.CENTER
    marca.add_run().italic = True

    document.add_heading('CARTEIRAS PERSONALIZADAS', 0)

    p = document.add_paragraph('Ago/2023')

    document.add_heading('As melhores ações com baixo custo e altos dividendos', level=1)
    document.add_paragraph('Modelo Magic Formula', style='Intense Quote')

    tabela = pd.read_csv('modelo_01.csv')

    n_rows = tabela.shape[0]
    n_cols = tabela.shape[1]

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Papel'
    hdr_cells[1].text = 'Cotação R$'
    hdr_cells[2].text = 'Cot.Máxima R$'
    hdr_cells[3].text = 'Div.Yield %'

    table = document.add_table(rows=n_rows, cols=n_cols)

    for col_i, col_l in enumerate(tabela):
        for row_i in range(n_rows):
            cell_ = table.cell(row_i, col_i)
            cell_.text = str(tabela[col_l][row_i])

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    document.add_page_break()

    document.add_heading('Muito Obrigado:\n', level=1)
    p = document.add_paragraph('Esta carteira de ações foi elaborada automaticamente por um programa em Python utilizando como critério de escolha os indicadores apontados por Joel Greenblatt em seu best seller “A Fórmula Mágica de Joel Greenblatt para Bater o Mercado de Ações” e os critérios sobre preço expostos por Louise Barsi em um podcast gravado com Tiago Nigro.\n')
    p = document.add_paragraph('A Fórmula Mágica de Joel Greenblatt é um modelo de investimentos de longo prazo com rentabilidade astronômica comprovada por mais de 15 anos de aplicações feitas pelo autor. Louise Barsi hoje é uma das referências nacionais mais bem conceituadas no mundo dos investimentos. Tiago Nigro é dono de um dos maiores canais nacionais de investimentos do youtube com uma grande carreira no mundo dos investimentos.\n')
    p = document.add_paragraph('Ressaltamos, com tudo, que este projeto tem como pretensão principal ser de cunho didático não sendo com isso de nenhuma maneira uma recomendação de investimento resguardando aos envolvidos de quaisquer más interpretações a respeito.\n')
    document.save('Carteiras/modelo_01.docx')
# Do_docx()